import streamlit as st
import pandas as pd
import requests
import time
import json
import os
import threading
from datetime import datetime
from bs4 import BeautifulSoup
from docx import Document
from io import BytesIO

# --- PAGE CONFIG ---
st.set_page_config(page_title="SEO & GEO Content Engineer", layout="wide", page_icon="✍️")

# --- CONSTANTS ---
HISTORY_FILE = "article_history.json"
# Hide from Streamlit watcher to avoid inotify limit
RUNNING_DIR = ".running_tasks"
if not os.path.exists(RUNNING_DIR):
    os.makedirs(RUNNING_DIR)

# --- CUSTOM CSS (Premium Dark Mode & Glassmorphism) ---
st.markdown("""
    <style>
    .main {
        background: linear-gradient(135deg, #0f0c29, #302b63, #24243e);
        color: #ffffff;
    }
    .stApp {
        background: transparent;
    }
    [data-testid="stSidebar"] {
        background: rgba(255, 255, 255, 0.05);
        backdrop-filter: blur(10px);
        border-right: 1px solid rgba(255, 255, 255, 0.1);
    }
    .stButton>button {
        background: linear-gradient(45deg, #00c6ff, #0072ff);
        color: white;
        border: none;
        border-radius: 8px;
        padding: 0.6rem 1.2rem;
        transition: all 0.3s ease;
    }
    .stButton>button:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 15px rgba(0, 198, 255, 0.4);
    }
    .status-box {
        background: rgba(255, 255, 255, 0.03);
        border: 1px solid rgba(255, 255, 255, 0.1);
        border-radius: 12px;
        padding: 1.5rem;
        margin-bottom: 1rem;
    }
    h1, h2, h3 {
        background: linear-gradient(to right, #00c6ff, #0072ff);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    </style>
""", unsafe_allow_html=True)

# --- INITIALIZATION ---
if 'api_keys' not in st.session_state:
    st.session_state.api_keys = {
        'AI': '',
        'SERP': '',
        'Linkup': ''
    }

if 'articles' not in st.session_state:
    st.session_state.articles = {}

# --- SIDEBAR CONFIG ---
with st.sidebar:
    st.title("⚙️ Cấu hình")
    st.session_state.api_keys['AI'] = st.text_input("AI API Key (chiasegpu)", type="password")
    st.session_state.api_keys['SERP'] = st.text_input("Serp API Key (serper.dev)", type="password")
    st.session_state.api_keys['Linkup'] = st.text_input("Linkup API Key", type="password")
    
    st.divider()
    st.info("💡 Mẹo: Nhập đầy đủ API Key để bắt đầu chạy workflow.")

# --- NAVIGATION TABS ---
tab1, tab2, tab3 = st.tabs(["🚀 Writer", "📜 Rules Editor", "📁 History"])

# --- TAB 2: RULES EDITOR ---
with tab2:
    st.header("Chỉnh sửa Quy tắc SEO & GEO")
    rules_path = "seo_geo_rules.md"
    if os.path.exists(rules_path):
        with open(rules_path, "r", encoding="utf-8") as f:
            rules_content = f.read()
    else:
        rules_content = "# SEO & GEO Rules\n\n(Tạo quy tắc của bạn tại đây)"
    
    updated_rules = st.text_area("Nội dung Markdown", value=rules_content, height=500)
    if st.button("Lưu quy tắc"):
        with open(rules_path, "w", encoding="utf-8") as f:
            f.write(updated_rules)
        st.success("Đã cập nhật quy tắc!")

# --- CORE FUNCTIONS ---

def truncate_text(text, max_chars=4000):
    if not isinstance(text, str):
        text = str(text)
    if len(text) > max_chars:
        return text[:max_chars] + "\n...[Nội dung đã bị cắt bớt để bảo toàn token]..."
    return text

def extract_json(text):
    try:
        # Xử lý Markdown code blocks
        if "```json" in text:
            text = text.split("```json")[1].split("```")[0]
        elif "```" in text:
            text = text.split("```")[1].split("```")[0]
        
        # Tìm dấu ngoặc nhọn đầu và cuối
        start = text.find('{')
        end = text.rfind('}')
        if start != -1 and end != -1:
            return text[start:end+1]
        return text.strip()
    except:
        return text.strip()

def repair_json(broken_json):
    if not broken_json:
        return ""
    stack = []
    in_string = False
    escaped = False
    
    # Dọn dẹp sơ bộ
    broken_json = broken_json.strip()
    
    for char in broken_json:
        if char == '\\' and in_string:
            escaped = not escaped
            continue
        if char == '"' and not escaped:
            in_string = not in_string
        if not in_string:
            if char in '{[':
                stack.append(char)
            elif char in '}]':
                if stack:
                    if (char == '}' and stack[-1] == '{') or (char == ']' and stack[-1] == '['):
                        stack.pop()
        escaped = False
    
    if in_string:
        broken_json += '"'
    
    while stack:
        opener = stack.pop()
        broken_json += '}' if opener == '{' else ']'
    return broken_json

def clean_ai_html(text):
    """Remove markdown code blocks like ```html or ``` from AI output."""
    text = text.strip()
    if text.startswith("```"):
        lines = text.splitlines()
        if lines[0].startswith("```"):
            lines = lines[1:]
        if lines and lines[-1].startswith("```"):
            lines = lines[:-1]
        return "\n".join(lines).strip()
    return text

def export_to_docx(html_content, keyword, meta_title="", meta_description=""):
    """Convert HTML structure to professional Word document."""
    doc = Document()
    doc.add_heading(meta_title if meta_title else keyword, 0)
    if meta_description:
        doc.add_paragraph(f"SEO Meta Description: {meta_description}")
        doc.add_paragraph("-" * 30)
    
    soup = BeautifulSoup(html_content, "html.parser")
    # Duyệt qua các thẻ con trực tiếp của body hoặc tệp tin
    for tag in soup.find_all(True):
        if tag.name == 'h2':
            doc.add_heading(tag.get_text(), level=1)
        elif tag.name == 'h3':
            doc.add_heading(tag.get_text(), level=2)
        elif tag.name == 'p':
            if tag.parent.name not in ['li']: # Tránh lặp lại thẻ trong danh sách
                doc.add_paragraph(tag.get_text())
        elif tag.name == 'li':
            doc.add_paragraph(tag.get_text(), style='List Bullet')
    
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

HISTORY_FILE = "article_history.json"

def load_history():
    if os.path.exists(HISTORY_FILE):
        with open(HISTORY_FILE, "r", encoding="utf-8") as f:
            try:
                return json.load(f)
            except:
                return []
    return []

def save_to_history(article):
    history = load_history()
    article['created_at'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    history.append(article)
    with open(HISTORY_FILE, "w", encoding="utf-8") as f:
        json.dump(history, f, ensure_ascii=False, indent=4)

def get_serp_results(keyword, api_key):
    url = "https://google.serper.dev/search"
    headers = {"X-API-KEY": api_key, "Content-Type": "application/json"}
    data = json.dumps({"q": keyword, "num": 10})
    response = requests.post(url, headers=headers, data=data)
    return response.json().get('organic', [])

def scrape_url(url):
    try:
        response = requests.get(url, timeout=10)
        soup = BeautifulSoup(response.text, 'html.parser')
        # Extract from main/article or fallback to body
        content = soup.find('main') or soup.find('article') or soup.find('body')
        return content.get_text(separator=' ', strip=True)[:5000] # Limit to 5000 chars
    except Exception as e:
        return f"Error scraping {url}: {str(e)}"

def call_ai(prompt, api_key, system_prompt="You are an expert SEO Content Engineer."):
    url = "https://llm.chiasegpu.vn/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    data = {
        "model": "claude-sonnet-4.6",
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": prompt}
        ],
        "max_tokens": 4000
    }
    try:
        response = requests.post(url, headers=headers, json=data, timeout=90)
        if response.status_code != 200:
            st.error(f"⚠️ API Error ({response.status_code}): {response.text[:300]}")
            return f"Error: {response.text[:200]}"
        
        return response.json()['choices'][0]['message']['content']
    except requests.exceptions.Timeout:
        st.error("⚠️ Lỗi: Thời gian phản hồi từ AI quá lâu (Timeout).")
        return "Error: Timeout"
    except requests.exceptions.HTTPError as e:
        if response.status_code in [502, 504]:
            st.error("⚠️ Lỗi Gateway (502/504): Có thể Payload quá lớn hoặc Server AI đang bảo trì.")
        return f"Error: {str(e)}"
    except Exception as e:
        st.error(f"⚠️ Lỗi hệ thống: {str(e)}")
        return f"Error: {str(e)}"

def call_ai_stream(prompt, api_key, system_prompt="You are an expert SEO Content Engineer."):
    """Generator to stream AI responses. Securely handles API keys for background threads."""
    url = "https://llm.chiasegpu.vn/v1/chat/completions"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    data = {"model": "claude-sonnet-4.6", "messages": [{"role": "system", "content": system_prompt}, {"role": "user", "content": prompt}], "stream": True, "max_tokens": 4000}
    response = None
    try:
        response = requests.post(url, headers=headers, json=data, timeout=90, stream=True)
        if response.status_code != 200:
            yield f"Error: {response.text[:200]}"
            return
        for line in response.iter_lines():
            if line:
                decoded_line = line.decode('utf-8').strip()
                if decoded_line.startswith("data: "):
                    content = decoded_line[6:]
                    if content == "[DONE]": break
                    try:
                        json_data = json.loads(content)
                        chunk = json_data['choices'][0]['delta'].get('content')
                        if chunk: yield chunk
                    except: continue
    except GeneratorExit: return 
    except Exception as e: yield f" [Error: {str(e)}]"
    finally:
        if response:
            try: response.close()
            except: pass

def update_task_status(kw, status_data):
    """Save current task progress to disk."""
    path = os.path.join(RUNNING_DIR, f"{kw}.json")
    with open(path, "w", encoding="utf-8") as f:
        json.dump(status_data, f, ensure_ascii=False, indent=4)

def load_task_status(kw):
    """Load current task progress from disk with error handling for race conditions."""
    path = os.path.join(RUNNING_DIR, f"{kw}.json")
    if os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                content = f.read().strip()
                if not content:
                    return None
                return json.loads(content)
        except (json.JSONDecodeError, Exception):
            # File might be empty or being written by the background thread
            return None
    return None

def background_worker(kw, api_keys, mode, serp_manual="", linkup_manual="", rules_path="seo_geo_rules.md"):
    """Threaded function to perform full generation without st. commands."""
    try:
        status_data = {
            "keyword": kw, "status": "processing", "log": "Bắt đầu...", 
            "research": {}, "meta_title": "", "meta_description": "", "content": ""
        }
        update_task_status(kw, status_data)

        # 1. Translate Query
        status_data["log"] = "Đang quốc tế hóa truy vấn (English)..."
        update_task_status(kw, status_data)
        eng_p = f"Translate the keyword '{kw}' into a precise scientific research query in English. Output ONLY the query string."
        eng_query = ""
        for chunk in call_ai_stream(eng_p, api_keys['AI']): eng_query += chunk
        eng_query = clean_ai_html(eng_query.strip()) or kw

        # 2. Research
        if mode == "🚀 Tự động (Full Workflow)":
            status_data["log"] = "Đang nghiên cứu (Serp & Linkup)..."
            update_task_status(kw, status_data)
            
            # Serp
            serp_results = get_serp_results(kw, api_keys['SERP'])
            urls = [r['link'] for r in serp_results[:3]]
            
            # Scrape
            comp_content = ""
            for url in urls:
                scraped = scrape_url(url)
                comp_content += f"\n--- Source: {url} ---\n" + truncate_text(scraped, 500)
            comp_content = truncate_text(comp_content, 1500)
            
            # Linkup
            linkup_json = linkup_research(eng_query, api_keys['Linkup'])
            status_data["research"] = {
                "answer": linkup_json.get('answer', ''), 
                "urls": [s.get('url') for s in linkup_json.get('sources', [])],
                "sources": linkup_json.get('sources', [])
            }
        else:
            status_data["log"] = "Dữ liệu nghiên cứu thủ công."
            status_data["research"] = {"answer": "Manual", "urls": []}
            comp_content = serp_manual

        # 3. Outline
        status_data["log"] = "Đang lập dàn ý (Outline)..."
        update_task_status(kw, status_data)
        with open(rules_path, "r", encoding="utf-8") as f:
            raw_rules = f.read()
        
        # Determine language first
        lang_p = f"Based on the keyword '{kw}', identify its language. Translate English research data into this language for the outline. Output JSON."
        
        outline_p = f"""SEO RULES: {truncate_text(raw_rules, 3000)}
KEYWORD: {kw}
RESEARCH DATA: {status_data['research']['answer']}

TASK: Create a professional SEO content outline.
1. Identify keyword language.
2. Structure: Sapo, AI Overview Box, Key Takeaways, Body Headings (4-6), FAQ, References.
3. Output strictly valid JSON:
{{
  "language": "...",
  "meta_title": "...", 
  "meta_description": "...",
  "headings": [{"title": "...", "points": "..."}],
  "faq": [{"q": "...", "a": "..."}]
}}"""
        outline_res = ""
        for chunk in call_ai_stream(outline_p, api_keys['AI']): outline_res += chunk
        try:
            outline_data = json.loads(repair_json(extract_json(outline_res)))
            status_data["meta_title"] = outline_data.get('meta_title', '')
            status_data["meta_description"] = outline_data.get('meta_description', '')
            target_lang = outline_data.get('language', 'Vietnamese')
            update_task_status(kw, status_data)
        except:
            status_data["log"] = "Lỗi tạo dàn ý. Kiểm tra API."
            update_task_status(kw, status_data)
            return

        # 4. Sequential Writing (Strict Structure)
        research_context = f"Research Summary: {status_data['research']['answer']}\nSource List: {status_data['research']['urls']}"
        lang_instr = f"CRITICAL: Write everything in {target_lang}. Translate all English concepts naturally."

        # Segment definitions
        segments = [
            ("Sapo", f"Task: Write a high-intent Sapo (3-4 sentences) including a direct hook answers for the keyword. Rules: No <a> tags. {lang_instr}\nContext: {research_context}"),
            ("AI Overview", f"Task: Write a compact 'AI Overview' summary (50-80 words). {lang_instr}\nContext: {research_context}"),
            ("Key Takeaways", f"Task: Summarize 3-5 key scientific takeaways in a bulleted list. {lang_instr}\nContext: {research_context}"),
        ]

        # Process Segments
        for name, p in segments:
            status_data["log"] = f"Đang viết: {name}"
            update_task_status(kw, status_data)
            seg_out = ""
            for chunk in call_ai_stream(p, api_keys['AI']): seg_out += chunk
            clean_out = clean_ai_html(seg_out)
            
            if name == "AI Overview":
                clean_out = f'<div style="border: 2px solid #00c6ff; padding: 15px; border-radius: 10px; background: rgba(0, 198, 255, 0.05); margin: 20px 0;"><strong>🤖 AI Overview:</strong><br>{clean_out}</div>'
            elif name == "Key Takeaways":
                clean_out = f"<h3>Key Takeaways</h3>\n{clean_out}"
            
            status_data["content"] += f"\n\n{clean_out}"
            update_task_status(kw, status_data)

        # Body Headings (One by one)
        for i, h in enumerate(outline_data.get('headings', [])):
            status_data["log"] = f"Đang viết Heading {i+1}: {h['title']}"
            update_task_status(kw, status_data)
            
            h_p = f"""Task: Write a deep-dive SEO section for heading '{h['title']}'.
Points to cover: {h['points']}
Context: {research_context}
Language: {target_lang}
Rules: RAW HTML ONLY. Use <h4> for points if needed. STRICT: NO <a> tags. Cite sources as PLAIN TEXT (e.g. According to NCBI)."""
            
            h_out = f"<h2>{h['title']}</h2>\n"
            for chunk in call_ai_stream(h_p, api_keys['AI']): h_out += chunk
            status_data["content"] += f"\n\n{clean_ai_html(h_out)}"
            update_task_status(kw, status_data)

        # FAQ
        status_data["log"] = "Đang viết FAQ..."
        update_task_status(kw, status_data)
        faq_p = f"Task: Write a structured FAQ section based on research. Use <h3> for questions. {lang_instr}\nQuestions: {outline_data.get('faq')}"
        faq_out = "<h2>Frequently Asked Questions</h2>\n"
        for chunk in call_ai_stream(faq_p, api_keys['AI']): faq_out += chunk
        status_data["content"] += f"\n\n{clean_ai_html(faq_out)}"
        update_task_status(kw, status_data)

        # References
        status_data["log"] = "Đang tạo danh mục nguồn tham khảo..."
        update_task_status(kw, status_data)
        ref_p = f"Task: Create a 'References' section. List 3-5 authoritative source URLs from the following research as <a> links. {lang_instr}\nData: {research_context}"
        ref_out = "<h2>References</h2>\n"
        for chunk in call_ai_stream(ref_p, api_keys['AI']): ref_out += chunk
        status_data["content"] += f"\n\n{clean_ai_html(ref_out)}"
        
        # Complete
        status_data["status"] = "complete"
        status_data["log"] = "✅ Đã xong!"
        update_task_status(kw, status_data)
        
        # Final save to history
        save_to_history({
            "keyword": kw, "meta_title": status_data["meta_title"],
            "meta_description": status_data["meta_description"], "content": status_data["content"]
        })
        # Cleanup
        os.remove(os.path.join(RUNNING_DIR, f"{kw}.json"))

    except Exception as e:
        status_data["status"] = "error"
        status_data["log"] = f"Lỗi: {str(e)}"
        update_task_status(kw, status_data)

def linkup_research(keyword, api_key):
    url = "https://api.linkup.so/v1/search"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    data = {
        "q": f"scientific evidence for {keyword}",
        "depth": "standard",
        "outputType": "sourcedAnswer"
    }
    response = requests.post(url, headers=headers, json=data)
    return response.json()

# --- TAB 1: WRITER ---
with tab1:
    st.header("Trình tạo nội dung SEO")
    
    mode = st.radio("Chế độ vận hành:", ["🚀 Tự động (Full Workflow)", "✍️ Thủ công (Dán dữ liệu)"], horizontal=True)

    keywords = []
    
    if mode == "🚀 Tự động (Full Workflow)":
        col1, col2 = st.columns([1, 1])
        with col1:
            manual_keywords = st.text_area("Nhập danh sách từ khóa (mỗi dòng 1 từ)", height=150)
        with col2:
            uploaded_file = st.file_uploader("Hoặc tải lên file CSV/XLSX", type=["csv", "xlsx"])

        if manual_keywords:
            keywords = [k.strip() for k in manual_keywords.split('\n') if k.strip()]
        if uploaded_file:
            if uploaded_file.name.endswith('.csv'):
                df = pd.read_csv(uploaded_file)
            else:
                df = pd.read_excel(uploaded_file)
            target_col = next((c for c in df.columns if c.lower() in ['keyword', 'từ khóa', 'tu khoa']), df.columns[0])
            keywords.extend(df[target_col].tolist())
    else:
        # MANUAL MODE INPUTS
        kw_manual = st.text_input("Từ khóa chính (Keyword)")
        col_m1, col_m2 = st.columns(2)
        with col_m1:
            research_manual = st.text_area("Dữ liệu nghiên cứu (Linkup/Scientific Data)", height=250, placeholder="Dán dữ liệu nghiên cứu khoa học tại đây...")
        with col_m2:
            competitor_manual = st.text_area("Nội dung đối thủ (Competitor Content)", height=250, placeholder="Dán nội dung đối thủ hoặc ý tưởng tại đây...")
        
        if kw_manual:
            keywords = [kw_manual.strip()]

    if st.button("🚀 Bắt đầu Viết bài"):
        if not all(st.session_state.api_keys.values()):
            st.error("Vui lòng điền đầy đủ API Key trong Sidebar.")
        elif not keywords:
            st.warning("Vui lòng nhập từ khóa.")
        else:
            for kw in keywords:
                # Launch thread if not already running
                task_path = os.path.join(RUNNING_DIR, f"{kw}.json")
                if not os.path.exists(task_path):
                    # Prepare manual data if in manual mode
                    c_man = competitor_manual if mode == "✍️ Thủ công (Dán dữ liệu)" else ""
                    r_man = research_manual if mode == "✍️ Thủ công (Dán dữ liệu)" else ""
                    
                    t = threading.Thread(
                        target=background_worker, 
                        args=(kw, st.session_state.api_keys, mode, c_man, r_man)
                    )
                    t.start()
            
            st.success(f"🚀 Đã bắt đầu xử lý {len(keywords)} bài viết chạy ngầm. Bạn có thể đóng tab, khóa máy hoặc chuyển trình duyệt!")
    # --- DYNAMIC MONITORING UI ---
    running_files = []
    if os.path.exists(RUNNING_DIR):
        running_files = [f for f in os.listdir(RUNNING_DIR) if f.endswith(".json")]
    
    if running_files:
        st.divider()
        st.subheader("⚙️ Đang xử lý ngầm...")
        for f in running_files:
            kwname = f.replace(".json", "")
            t_status = load_task_status(kwname)
            if t_status:
                with st.expander(f"🔄 {kwname}: {t_status['log']}", expanded=True):
                    # Progress estimate
                    prog = 0.1
                    l_lower = t_status['log'].lower()
                    if "nghiên cứu" in l_lower: prog = 0.3
                    if "dàn ý" in l_lower: prog = 0.5
                    if "đang viết" in l_lower: prog = 0.8
                    if "✅" in t_status['log']: prog = 1.0
                    st.progress(prog)

                    # SHOW RESEARCH IF AVAILABLE
                    if t_status.get('research') and t_status['research'].get('answer'):
                        with st.expander("🔬 Dữ liệu nghiên cứu (Research Results)"):
                            st.write(t_status['research']['answer'])
                            if t_status['research'].get('urls'):
                                st.markdown("**🔗 Nguồn tham khảo:**")
                                for u in t_status['research']['urls']:
                                    st.write(f"- {u}")

                    # SHOW OUTLINE IF AVAILABLE
                    if t_status.get('meta_title'):
                        with st.expander("📝 Dàn ý SEO (Outline)"):
                            st.success(f"**Meta Title:** {t_status['meta_title']}")
                            st.info(f"**Meta Description:** {t_status['meta_description']}")

                    if t_status['content']:
                        st.info("💡 Nội dung bài viết (Live Preview):")
                        # Show content in a scrollable area
                        st.markdown(f'<div style="max-height: 400px; overflow-y: auto; border: 1px solid #ddd; padding: 15px; border-radius: 10px;">{t_status["content"]}</div>', unsafe_allow_html=True)
        
        if st.button("🔄 Làm mới thủ công"):
            st.rerun()
        
        time.sleep(10)
        st.rerun()

    # --- DISPLAY PERSISTENT RESULTS ---
    if st.session_state.articles:
        st.divider()
        st.header("📝 Kết quả đã tạo")
        if st.button("🧹 Xóa tất cả kết quả"):
            st.session_state.articles = {}
            st.rerun()

        for kw, data in st.session_state.articles.items():
            with st.expander(f"📄 Bài viết: {kw} ({data['status']})", expanded=(data['status'] == "complete")):
                if data['research']:
                    st.info(f"📊 Debug Metrics: {data['metrics']}")
                    with st.expander("🔍 Chi tiết nghiên cứu"):
                        st.write(data['research'].get('answer', "N/A"))
                        if 'raw' in data['research']:
                            st.json(data['research']['raw'])
                
                if data['meta_title']:
                    st.success(f"**SEO Meta Title:** {data['meta_title']}")
                    st.info(f"**SEO Meta Description:** {data['meta_description']}")
                
                st.markdown("### Nội dung bài viết")
                st.html(data['content'])
                
                if data['status'] == "complete":
                    col_ex1, col_ex2 = st.columns(2)
                    # Word Export Professional
                    docx_data = export_to_docx(data['content'], kw, data['meta_title'], data['meta_description'])
                    col_ex1.download_button(label=f"📥 Tải .docx ({kw})", data=docx_data, file_name=f"{kw}.docx")
                    
                    # HTML Export Clean
                    html_full = f"<html><head><meta charset='utf-8'><title>{data['meta_title']}</title><meta name='description' content='{data['meta_description']}'></head><body>{data['content']}</body></html>"
                    col_ex2.download_button(label=f"📥 Tải .html ({kw})", data=html_full, file_name=f"{kw}.html")

# --- TAB 3: HISTORY ---
with tab3:
    st.header("🕒 Lịch sử bài viết vĩnh viễn")
    history_data = load_history()
    
    if not history_data:
        st.info("Chưa có bài viết nào trong lịch sử.")
    else:
        st.write(f"Đang có {len(history_data)} bài viết đã được lưu trữ.")
        
        # Clear History Logic
        with st.expander("🛠️ Quản lý kho lưu trữ", expanded=False):
            confirm = st.checkbox("Tôi xác nhận muốn xóa vĩnh viễn toàn bộ lịch sử")
            if st.button("🗑️ Xóa sạch vĩnh viễn", disabled=not confirm):
                if os.path.exists(HISTORY_FILE):
                    os.remove(HISTORY_FILE)
                st.success("Đã xóa sạch lịch sử!")
                st.rerun()

        st.divider()
        
        # Display history items in reverse (newest first)
        for item in reversed(history_data):
            kw_name = item.get('keyword', 'N/A')
            created_at = item.get('created_at', 'N/A')
            with st.expander(f"📅 {created_at} | 🔑 {kw_name}"):
                st.success(f"**Title:** {item.get('meta_title', 'N/A')}")
                st.markdown(item.get('content', ''), unsafe_allow_html=True)
                
                col1, col2 = st.columns(2)
                # Word Export
                docx_hist = export_to_docx(item.get('content', ''), kw_name, item.get('meta_title', ''), item.get('meta_description', ''))
                col1.download_button(label="📥 Tải .docx", data=docx_hist, file_name=f"{kw_name}.docx", key=f"dl_docx_{created_at}")
                # HTML Export
                html_hist = f"<html><head><meta charset='utf-8'><title>{item.get('meta_title')}</title></head><body>{item.get('content')}</body></html>"
                col2.download_button(label="📥 Tải .html", data=html_hist, file_name=f"{kw_name}.html", key=f"dl_html_{created_at}")
